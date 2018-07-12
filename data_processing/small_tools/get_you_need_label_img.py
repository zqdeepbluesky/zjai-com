# -*- coding: utf-8 -*-
# @Time    : 5/24/2018
# @Author  : CarrieChen
# @File    : get_you_need_label_img.py
# @Software: ZJ_AI
# this code is for read some labels from excel and  find according imgs and put the imgs into a word.

import xlrd
import docx
from PIL import Image
import os
import xlwt
import io_utils
from PIL import ImageDraw
from PIL import ImageFont

#some paths
parent_path="C:\\Users\\Administrator\\Desktop\\data_processing_carriechen\\count_all_annotations"
excel_path=parent_path+"\\本批商品列表.xls"
img_path=parent_path+"\\pic+lab166"
refer_166classes=parent_path+"\\166_classes_list.xls"
this_batch_imgs_path=parent_path+"\\本批商品图例"


def get_labels(input_path):
    data=xlrd.open_workbook(input_path)
    table=data.sheets()[0]
    labels=table.col_values(0)
    return labels

def get_chinese(input_path,pointlabel): #excel
    data=xlrd.open_workbook(input_path)
    table=data.sheets()[0]
    labels=table.col_values(4)
    if pointlabel in labels:  #else
        row=labels.index(pointlabel)
        product=table.cell(row,0).value
        taste = table.cell(row, 1).value
        weight=table.cell(row,2).value
        package=table.cell(row,3).value
    return product,taste,weight,package


def find_imgs_and_write_word(labels,parent_path):
    file=docx.Document()
    for i in range(len(labels)):
        img=img_path+"\\"+labels[i]+".jpg"
        product,taste,weight,package=get_chinese(refer_166classes,labels[i])
        file.add_picture(img)
        file.add_paragraph(product+taste+weight+package+" "+labels[i])
        file.add_paragraph("\n")
    file.save(parent_path+"\\"+"本批商品图例.doc")  #这是生成的word文档的名字


def find_imgs_and_save_as_imgs(labels, parent_path):
    io_utils.mkdir(this_batch_imgs_path)
    for i in range(len(labels)):
        background = Image.open("C:\\Users\\Administrator\\Desktop\\data_processing_carriechen\\count_all_annotations\\pure_white_background.jpg")
        img = img_path + "\\" + labels[i] + ".jpg"
        product, taste, weight, package = get_chinese(refer_166classes, labels[i])
        img=Image.open(img)
        background.paste(img,[100,50])
        draw = ImageDraw.Draw(background)
        width, height = background.size
        setFont = ImageFont.truetype('C:\Windows\Fonts\\simfang.ttf', 30)
        fillColor = "black"
        draw.text((10, height - 100), u"\""+labels[i]+"\"", font=setFont, fill=fillColor)
        draw.text((10, height - 50), u"\"" + product+taste+weight+package + "\"", font=setFont, fill=fillColor)
        background.save(this_batch_imgs_path+"\\"+labels[i]+".jpg")

if __name__ =="__main__":
    labels=get_labels(excel_path)
    #find_imgs_and_write_word(labels,parent_path)
    find_imgs_and_save_as_imgs(labels,parent_path)





